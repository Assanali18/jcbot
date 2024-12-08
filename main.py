from aiogram import Bot, Dispatcher, types
from aiogram.filters import Command, StateFilter
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from docx import Document
from docx.shared import Pt
import base64
import requests
import asyncio
import logging
import os
from dotenv import load_dotenv

load_dotenv()

ACCESS_TOKEN = os.getenv('ACCESS_TOKEN')
CLIENT_ID = os.getenv("CLIENT_ID")
CLIENT_SECRET = os.getenv("CLIENT_SECRET")
TENANT_ID = os.getenv("TENANT_ID")
REDIRECT_URI = os.getenv("REDIRECT_URI")
TOKEN_URL = f"https://login.microsoftonline.com/{TENANT_ID}/oauth2/v2.0/token"
REFRESH_TOKEN = os.getenv("REFRESH_TOKEN")


TELEGRAM_TOKEN = os.getenv("TELEGRAM_TOKEN")
ALLOWED_USER_ID = int(os.getenv("ALLOWED_USER_ID"))


TEST_RECIPIENTS = os.getenv("TEST_RECIPIENTS").split(",")
PROD_RECIPIENTS = os.getenv("PROD_RECIPIENTS").split(",")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

bot = Bot(token=TELEGRAM_TOKEN)
dp = Dispatcher(storage=MemoryStorage())


class Form(StatesGroup):
    departure_date = State()
    arrival_date = State()


current_mode = "test"


def refresh_access_token():
    global REFRESH_TOKEN
    payload = {
        "client_id": CLIENT_ID,
        "client_secret": CLIENT_SECRET,
        "grant_type": "refresh_token",
        "refresh_token": REFRESH_TOKEN,
        "redirect_uri": REDIRECT_URI,
    }
    response = requests.post(TOKEN_URL, data=payload)
    if response.status_code == 200:
        tokens = response.json()
        REFRESH_TOKEN = tokens.get("refresh_token", REFRESH_TOKEN)
        logger.info("Токен успешно обновлен.")
        return tokens["access_token"]
    else:
        logger.error(f"Ошибка обновления токена: {response.text}")
        raise Exception("Не удалось обновить токен.")


async def update_status(message, status_message, new_text):
    try:
        await bot.edit_message_text(
            chat_id=message.chat.id, message_id=status_message.message_id, text=new_text
        )
    except Exception as e:
        logger.error(f"Ошибка обновления сообщения: {e}")


def fill_document(departure_date, arrival_date):
    logger.info("Заполнение документа...")
    doc = Document("template.docx")
    for paragraph in doc.paragraphs:
        if "«_____» ________" in paragraph.text:
            paragraph.text = paragraph.text.replace(
                "«_____» ________", f"«{departure_date}»"
            ).replace(
                "«_____»  ____________", f"«{arrival_date}»"
            )

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)

    output_file = "filled_document.docx"
    doc.save(output_file)
    logger.info("Документ успешно заполнен.")
    return output_file


def send_email(file_path, recipients):
    global ACCESS_TOKEN
    logger.info(f"Отправка письма. Получатели: {recipients}")
    with open(file_path, "rb") as file:
        encoded_file = base64.b64encode(file.read()).decode('utf-8')

    url = "https://graph.microsoft.com/v1.0/me/sendMail"
    headers = {
        "Authorization": f"Bearer {ACCESS_TOKEN}",
        "Content-Type": "application/json"
    }

    email_data = {
        "message": {
            "subject": "Заявление на временное покидание",
            "body": {
                "contentType": "Text",
                "content": "Пожалуйста, найдите заявление на временное покидание во вложении."
            },
            "toRecipients": [{"emailAddress": {"address": email}} for email in recipients],
            "attachments": [
                {
                    "@odata.type": "#microsoft.graph.fileAttachment",
                    "name": "Заявление.docx",
                    "contentBytes": encoded_file
                }
            ]
        },
        "saveToSentItems": "true"
    }

    response = requests.post(url, headers=headers, json=email_data)
    if response.status_code == 401 and "InvalidAuthenticationToken" in response.text:
        logger.warning("Токен истек. Обновление токена...")
        ACCESS_TOKEN = refresh_access_token()
        headers["Authorization"] = f"Bearer {ACCESS_TOKEN}"
        response = requests.post(url, headers=headers, json=email_data)

    logger.info(f"Результат отправки: {response.status_code}, {response.text}")
    return response.status_code, response.text


async def is_user_allowed(message: types.Message):
    if message.from_user.id != ALLOWED_USER_ID:
        await message.answer("У вас нет доступа к этому боту.")
        return False
    return True


@dp.message(Command("start"))
async def start(message: types.Message, state: FSMContext):
    if not await is_user_allowed(message):
        return

    await message.answer(f"Привет! Сейчас вы на {current_mode} режиме.Введите дату выезда (например, 15 октября):")
    await state.set_state(Form.departure_date)


@dp.message(Form.departure_date)
async def departure_date(message: types.Message, state: FSMContext):
    if not await is_user_allowed(message):
        return

    await state.update_data(departure_date=message.text)
    await message.answer("Введите дату заезда (например, 20 октября):")
    await state.set_state(Form.arrival_date)


@dp.message(Form.arrival_date)
async def arrival_date(message: types.Message, state: FSMContext):
    if not await is_user_allowed(message):
        return

    await state.update_data(arrival_date=message.text)
    data = await state.get_data()
    departure_date = data["departure_date"]
    arrival_date = data["arrival_date"]

    status_message = await message.answer("<b>Заполнение документа...</b>", parse_mode="HTML")
    file_path = fill_document(departure_date, arrival_date)
    await update_status(message, status_message, "Документ успешно заполнен.")

    recipients = TEST_RECIPIENTS if current_mode == "test" else PROD_RECIPIENTS
    status_message = await message.answer("<b>Отправка письма...</b>", parse_mode="HTML")
    status_code, response_text = send_email(file_path, recipients)

    if status_code == 202:
        await update_status(
            message,
            status_message,
            f"Письмо успешно отправлено! Режим: {current_mode.capitalize()}, получатели: {', '.join(recipients)} ",
        )
    else:
        await update_status(message, status_message, f"Ошибка отправки письма: {response_text}")

    await state.clear()


@dp.message(Command("mode"), ~StateFilter())
async def change_mode(message: types.Message, state: FSMContext):
    if not await is_user_allowed(message):
        return

    global current_mode
    current_mode = "prod" if current_mode == "test" else "test"
    await message.answer(f"Режим изменен на {current_mode.capitalize()}.")

    await state.clear()


async def main():
    await dp.start_polling(bot)


if __name__ == "__main__":
    logger.info("Запуск бота...")
    asyncio.run(main())
